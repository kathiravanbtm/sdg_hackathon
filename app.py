from flask import Flask, request, send_file, render_template
from docx import Document
import io
import os

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

@app.route('/')
def index():
    return render_template('index.html')  # Serves the HTML form

@app.route('/generate', methods=['POST'])
def generate_doc():
    template_path = os.path.join(BASE_DIR, "template.docx")

    if not os.path.exists(template_path):
        return "Error: template.docx not found!", 404

    doc = Document(template_path)

    # 🔹 1️⃣ Collect input data & normalize spaces
    data = {key.strip(): request.form[key].strip().replace("\n", " ") for key in request.form}

    # 🔹 2️⃣ Log input data for debugging
    validation_file_path = os.path.join(BASE_DIR, "form_data.txt")
    with open(validation_file_path, "w", encoding="utf-8") as f:
        for key, value in data.items():
            f.write(f"{key}: {value}\n")

    # 🔹 3️⃣ Process Dynamic Objectives
    objectives = request.form.getlist('objective')  # Get all dynamically added objectives

    # If no objectives are provided, add a default placeholder
    if not objectives:
        objectives.append("No objectives provided")

    objectives_text = "\n".join([f"{i+1}. {obj}" for i, obj in enumerate(objectives)])

    # Ensure `{Objectives}` is correctly replaced
    data["Objectives"] = objectives_text  

    def replace_text_in_runs(runs, key, value):
        """Replaces placeholders inside runs to maintain formatting."""
        for run in runs:
            if key in run.text:
                run.text = run.text.replace(key, value)

    # 🔹 4️⃣ Replace `{Objectives}` placeholder in all paragraphs
    for para in doc.paragraphs:
        replace_text_in_runs(para.runs, "{Objectives}", objectives_text)

    # 🔹 5️⃣ Replace placeholders inside tables (handles text inside table cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_runs(para.runs, "{Objectives}", objectives_text)

    # 🔹 6️⃣ Replace all remaining placeholders
    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{key}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"{{{key}}}"
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, value)

    # 🔹 7️⃣ Save the document to memory and return it
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(file_stream, as_attachment=True, download_name="Course_Syllabus.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    app.run(debug=True)
